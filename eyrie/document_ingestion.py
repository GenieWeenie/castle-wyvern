"""
Castle Wyvern Document Ingestion System (RAG)
Feature 9: Document ingestion with Retrieval Augmented Generation
"""

import os
import json
import hashlib
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
import re

# For text extraction
try:
    import PyPDF2

    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx

    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False


@dataclass
class DocumentChunk:
    """A chunk of a document with metadata."""

    id: str
    document_id: str
    content: str
    chunk_index: int
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None

    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class Document:
    """A stored document."""

    id: str
    filename: str
    file_path: str
    file_type: str
    content: str
    chunks: List[str]  # Chunk IDs
    metadata: Dict[str, Any]
    created_at: str

    def to_dict(self) -> Dict:
        return asdict(self)


class SimpleEmbedding:
    """
    Simple embedding generator using keyword-based approach.
    For production, replace with OpenAI embeddings or local model.
    """

    def __init__(self):
        # Common keywords for simple semantic matching
        self.vocabulary = self._build_vocabulary()

    def _build_vocabulary(self) -> Dict[str, int]:
        """Build a simple keyword vocabulary."""
        # Technology keywords
        tech_words = [
            "python",
            "javascript",
            "java",
            "rust",
            "go",
            "code",
            "function",
            "class",
            "api",
            "database",
            "sql",
            "server",
            "client",
            "http",
            "rest",
            "graphql",
            "security",
            "auth",
            "encrypt",
            "password",
            "token",
            "vulnerability",
            "architecture",
            "design",
            "pattern",
            "microservice",
            "monolith",
            "test",
            "testing",
            "pytest",
            "unittest",
            "mock",
            "deploy",
            "docker",
            "kubernetes",
            "ci/cd",
            "pipeline",
            "machine learning",
            "ai",
            "model",
            "training",
            "data",
            "frontend",
            "backend",
            "fullstack",
            "web",
            "mobile",
        ]

        # Business keywords
        business_words = [
            "user",
            "customer",
            "product",
            "feature",
            "requirement",
            "revenue",
            "cost",
            "profit",
            "market",
            "competitor",
            "team",
            "project",
            "milestone",
            "deadline",
            "sprint",
        ]

        all_words = tech_words + business_words
        return {word: i for i, word in enumerate(all_words)}

    def generate(self, text: str) -> List[float]:
        """
        Generate a simple embedding vector.
        This is a keyword-based approach - not as good as neural embeddings
        but works without external API calls.
        """
        text_lower = text.lower()
        vector = [0.0] * len(self.vocabulary)

        # Count keyword occurrences
        for word, idx in self.vocabulary.items():
            count = len(re.findall(r"\b" + word + r"\b", text_lower))
            vector[idx] = float(count)

        # Normalize
        total = sum(vector)
        if total > 0:
            vector = [v / total for v in vector]

        return vector

    def similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors."""
        if len(vec1) != len(vec2):
            return 0.0

        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        magnitude1 = sum(a * a for a in vec1) ** 0.5
        magnitude2 = sum(b * b for b in vec2) ** 0.5

        if magnitude1 == 0 or magnitude2 == 0:
            return 0.0

        return dot_product / (magnitude1 * magnitude2)


class DocumentIngestion:
    """
    Document ingestion and RAG (Retrieval Augmented Generation) system.

    Features:
    - Import documents (txt, md, py, pdf, docx)
    - Chunk documents for retrieval
    - Simple embeddings for semantic search
    - Query to find relevant document sections
    - Integration with Broadway for summarization
    """

    def __init__(self, storage_dir: str = "grimoorum/documents"):
        self.storage_dir = Path(storage_dir)
        self.docs_file = self.storage_dir / "documents.json"
        self.chunks_file = self.storage_dir / "chunks.json"

        self.documents: Dict[str, Document] = {}
        self.chunks: Dict[str, DocumentChunk] = {}
        self.embedding = SimpleEmbedding()

        self._initialize_storage()
        self._load_data()

    def _initialize_storage(self):
        """Create storage directory if needed."""
        self.storage_dir.mkdir(parents=True, exist_ok=True)

        for file_path in [self.docs_file, self.chunks_file]:
            if not file_path.exists():
                with open(file_path, "w") as f:
                    json.dump({}, f)

    def _load_data(self):
        """Load existing documents and chunks."""
        try:
            with open(self.docs_file, "r") as f:
                data = json.load(f)
                self.documents = {d["id"]: Document(**d) for d in data.values()}

            with open(self.chunks_file, "r") as f:
                data = json.load(f)
                self.chunks = {c["id"]: DocumentChunk(**c) for c in data.values()}
        except (FileNotFoundError, json.JSONDecodeError):
            self.documents = {}
            self.chunks = {}

    def _save_data(self):
        """Persist documents and chunks."""
        with open(self.docs_file, "w") as f:
            json.dump({k: asdict(v) for k, v in self.documents.items()}, f, indent=2)

        with open(self.chunks_file, "w") as f:
            json.dump({k: asdict(v) for k, v in self.chunks.items()}, f, indent=2)

    def ingest_file(self, file_path: str, metadata: Dict = None) -> str:
        """
        Ingest a document file.

        Supports: .txt, .md, .py, .js, .json, .pdf (if PyPDF2 installed), .docx (if python-docx)

        Returns:
            Document ID
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Extract content based on file type
        content = self._extract_content(file_path)

        # Generate document ID
        doc_id = hashlib.md5(f"{file_path}{datetime.now().isoformat()}".encode(), usedforsecurity=False).hexdigest()[:12]

        # Chunk the content
        chunks = self._chunk_content(content, doc_id, file_path.name)
        chunk_ids = [c.id for c in chunks]

        # Store chunks
        for chunk in chunks:
            self.chunks[chunk.id] = chunk

        # Create document record
        document = Document(
            id=doc_id,
            filename=file_path.name,
            file_path=str(file_path),
            file_type=file_path.suffix.lower(),
            content=(
                content[:1000] + "..." if len(content) > 1000 else content
            ),  # Truncate for storage
            chunks=chunk_ids,
            metadata=metadata or {},
            created_at=datetime.now().isoformat(),
        )

        self.documents[doc_id] = document
        self._save_data()

        return doc_id

    def _extract_content(self, file_path: Path) -> str:
        """Extract text content from various file types."""
        suffix = file_path.suffix.lower()

        # Text files
        if suffix in [".txt", ".md", ".py", ".js", ".json", ".yaml", ".yml", ".html", ".css"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        # PDF files
        if suffix == ".pdf":
            if not PDF_SUPPORT:
                raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

            text = ""
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text

        # Word documents
        if suffix == ".docx":
            if not DOCX_SUPPORT:
                raise ImportError("python-docx not installed. Run: pip install python-docx")

            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])

        raise ValueError(f"Unsupported file type: {suffix}")

    def _chunk_content(
        self,
        content: str,
        document_id: str,
        filename: str,
        chunk_size: int = 500,
        overlap: int = 50,
    ) -> List[DocumentChunk]:
        """
        Split content into overlapping chunks.

        Args:
            content: Full document text
            document_id: Parent document ID
            filename: Source filename
            chunk_size: Characters per chunk
            overlap: Overlap between chunks
        """
        chunks = []

        # Simple sentence-based chunking
        sentences = re.split(r"(?<=[.!?])\s+", content)

        current_chunk = []
        current_size = 0
        chunk_index = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            sentence_size = len(sentence)

            # Start new chunk if current would exceed size
            if current_size + sentence_size > chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)

                # Generate embedding
                embedding = self.embedding.generate(chunk_text)

                chunk_id = f"{document_id}_chunk_{chunk_index}"
                chunk = DocumentChunk(
                    id=chunk_id,
                    document_id=document_id,
                    content=chunk_text,
                    chunk_index=chunk_index,
                    metadata={"source": filename, "type": "sentence_chunk"},
                    embedding=embedding,
                )
                chunks.append(chunk)

                # Overlap: keep last sentence(s) for context
                current_chunk = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk
                current_size = sum(len(s) for s in current_chunk)
                chunk_index += 1

            current_chunk.append(sentence)
            current_size += sentence_size

        # Don't forget the last chunk
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            embedding = self.embedding.generate(chunk_text)

            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunk = DocumentChunk(
                id=chunk_id,
                document_id=document_id,
                content=chunk_text,
                chunk_index=chunk_index,
                metadata={"source": filename, "type": "sentence_chunk"},
                embedding=embedding,
            )
            chunks.append(chunk)

        return chunks

    def search(self, query: str, top_k: int = 5) -> List[Dict]:
        """
        Search for relevant document chunks.

        Uses simple embedding similarity. Replace with proper vector DB for production.
        """
        if not self.chunks:
            return []

        # Generate query embedding
        query_embedding = self.embedding.generate(query)

        # Score all chunks
        scored_chunks = []
        for chunk in self.chunks.values():
            if chunk.embedding:
                score = self.embedding.similarity(query_embedding, chunk.embedding)
                scored_chunks.append((score, chunk))

        # Sort by score (highest first)
        scored_chunks.sort(key=lambda x: x[0], reverse=True)

        # Return top K with metadata
        results = []
        for score, chunk in scored_chunks[:top_k]:
            doc = self.documents.get(chunk.document_id)
            results.append(
                {
                    "chunk_id": chunk.id,
                    "document_name": doc.filename if doc else "Unknown",
                    "content": chunk.content,
                    "score": round(score, 3),
                    "chunk_index": chunk.chunk_index,
                }
            )

        return results

    def query_with_context(self, query: str, top_k: int = 3) -> str:
        """
        Get context string for AI prompts.

        Returns relevant chunks formatted for inclusion in prompts.
        """
        results = self.search(query, top_k)

        if not results:
            return ""

        lines = ["## Relevant Document Context", ""]

        for i, result in enumerate(results, 1):
            lines.append(
                f"### Source {i}: {result['document_name']} (relevance: {result['score']})"
            )
            lines.append(result["content"][:500])  # Limit chunk size
            lines.append("")

        return "\n".join(lines)

    def list_documents(self) -> List[Dict]:
        """List all ingested documents."""
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "type": doc.file_type,
                "chunks": len(doc.chunks),
                "created": doc.created_at,
            }
            for doc in self.documents.values()
        ]

    def get_document(self, doc_id: str) -> Optional[Dict]:
        """Get document details."""
        if doc_id not in self.documents:
            return None

        doc = self.documents[doc_id]
        return {
            "id": doc.id,
            "filename": doc.filename,
            "file_path": doc.file_path,
            "file_type": doc.file_type,
            "chunks": len(doc.chunks),
            "metadata": doc.metadata,
            "created_at": doc.created_at,
        }

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its chunks."""
        if doc_id not in self.documents:
            return False

        doc = self.documents[doc_id]

        # Delete chunks
        for chunk_id in doc.chunks:
            if chunk_id in self.chunks:
                del self.chunks[chunk_id]

        # Delete document
        del self.documents[doc_id]
        self._save_data()

        return True

    def get_stats(self) -> Dict:
        """Get ingestion system statistics."""
        return {
            "total_documents": len(self.documents),
            "total_chunks": len(self.chunks),
            "storage_dir": str(self.storage_dir),
            "supported_types": [".txt", ".md", ".py", ".js", ".json", ".pdf", ".docx"],
        }


if __name__ == "__main__":
    # Test document ingestion
    print("📚 Testing Document Ingestion System")
    print("=" * 50)

    ingestion = DocumentIngestion()

    # Create a test document
    test_doc = """
# Python Best Practices

## Functions
Functions should be small and do one thing. Use descriptive names.

## Classes
Classes should encapsulate related data and behavior. Follow the single responsibility principle.

## Error Handling
Always handle errors gracefully. Use try-except blocks and provide meaningful error messages.

## Testing
Write tests for your code. Use pytest for a modern testing experience.
"""

    # Save test file
    test_path = "/tmp/test_python_guide.md"
    with open(test_path, "w") as f:
        f.write(test_doc)

    # Ingest it
    print(f"\n1. Ingesting {test_path}...")
    doc_id = ingestion.ingest_file(test_path, {"topic": "python", "author": "Castle Wyvern"})
    print(f"   Document ID: {doc_id}")

    # Search
    print("\n2. Searching for 'testing'...")
    results = ingestion.search("testing", top_k=3)
    for r in results:
        print(f"   - {r['document_name']} (score: {r['score']})")
        print(f"     {r['content'][:80]}...")

    print("\n3. Stats:")
    stats = ingestion.get_stats()
    print(f"   Documents: {stats['total_documents']}")
    print(f"   Chunks: {stats['total_chunks']}")

    print("\n✅ Document ingestion working!")
